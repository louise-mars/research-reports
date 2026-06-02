from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default font for Chinese
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Title
title = doc.add_heading('AI转型每周简报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('日期：2026年5月11日（周一） | 编制：诸葛亮 · 首席战略顾问')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Section 1: International Politics
h1 = doc.add_heading('一、国际政治重大事件', level=1)

items1 = [
    ('1. 特朗普称伊朗对美和平提议的回应"完全不可接受"', '中东局势持续紧张，伊朗核问题再次成为全球焦点，地缘政治风险上升，能源市场可能波动'),
    ('2. 俄罗斯在非洲影响力受挫', '叛乱武装将普京军队赶出关键非洲城镇（马里基达尔），俄在非地区影响力正在减弱'),
    ('3. 美国情报搜集飞行在古巴外海激增', '冷战思维回流，美古关系再度紧张'),
    ('4. 委内瑞拉后马杜罗时代过渡谈判', '卡塔尔斡旋下，委内瑞拉政治过渡谈判进行中'),
    ('5. 荷兰当局关停儿童色情网站', '继CNN调查后，荷兰当局关停涉儿童性虐待平台，全球网络治理趋严'),
]

for title_text, content in items1:
    p = doc.add_paragraph()
    run = p.add_run(f'■ {title_text}')
    run.bold = True
    p.add_run(f'\n{content}')

doc.add_paragraph()

# Section 2: AI
h2 = doc.add_heading('二、AI/人工智能重大进展', level=1)

items2 = [
    ('1. 世界模型（World Models）：AI理解物理世界的新路径', '当前AI仍不可靠，需教AI理解周围世界才能解决此问题，从"统计拟合"向"物理理解"的范式转变'),
    ('2. 军事AI"战争室"：指挥官咨询系统', '军方已用AI检测人类可能遗漏的信息，现在还想要一个供指挥官在战斗中咨询的建议引擎'),
    ('3. 人形机器人数据：训练AI的人类运动数据需求激增', '机器人公司需要大量关于人类手部和肢体移动方式的数据，数据成为AI发展核心资源'),
    ('4. AI生殖医学：IVF自动化革新', '自动化、AI和筛查技术正在改变生殖医学，AI+医疗赛道持续火热'),
    ('5. AI图像生成：ChatGPT Images 2 发布', 'OpenAI推出全新图像生成模型，多模态AI竞争加剧'),
    ('6. AI版权争议：YouTube创作者起诉亚马逊', 'AI训练数据版权问题法律战升级，将倒逼行业建立数据授权机制'),
    ('7. AI Pentagon合同：Google不顾员工反对签署', '科技-军事深度融合，AI伦理争议持续'),
    ('8. 意图式混沌测试：为"自信且错误"的AI设计', '自主AI系统在生产环境中可能"自信地犯错"，企业级AI运维/安全赛道兴起'),
]

for title_text, content in items2:
    p = doc.add_paragraph()
    run = p.add_run(f'■ {title_text}')
    run.bold = True
    p.add_run(f'\n{content}')

doc.add_paragraph()

# Section 3: Communications
h3 = doc.add_heading('三、通信行业重大进展', level=1)

items3 = [
    ('1. 美国6G频谱政策最新进展', 'NTIA启动Spectrum.gov，寻求6G频谱资源，Verizon备战世界杯5G，6G标准战升温'),
    ('2. 英国电信（BT）准备今夏推出5G网络切片服务', '5G商用深化，网络切片商业化落地'),
    ('3. 塔杆公司对6G频谱、AI和边缘计算乐观', '基础设施层持续获资本青睐'),
    ('4. SKT净利再次下滑，数据泄露影响持续', '运营商数据安全信任危机'),
    ('5. 印度Reliance探索LEO卫星业务', '新兴市场卫星互联网竞争加剧，Starlink面临挑战'),
    ('6. Vodafone将主权云外包给AWS', '运营商云化趋势持续，公有云整合加速'),
    ('7. Deutsche Telekom从爱立信云转向Mavenir', 'Open RAN势起，传统设备商份额受压'),
    ('8. WNBA与AWS多年合作', '体育产业AI升级，AI落地应用场景持续扩展'),
]

for title_text, content in items3:
    p = doc.add_paragraph()
    run = p.add_run(f'■ {title_text}')
    run.bold = True
    p.add_run(f'\n{content}')

doc.add_paragraph()

# Section 4: Analysis
h4 = doc.add_heading('四、深度分析与趋势预测', level=1)

analysis_items = [
    ('重点关注一：伊朗局势升级', '中东火药桶再度引爆，未来2-3个月内不排除军事摩擦可能，建议资产配置中纳入黄金和能源股对冲'),
    ('重点关注二：AI军事化应用', 'AI已从"辅助工具"升级为"决策参与者"，2026-2027年将出现更多"AI指挥官"商业化产品'),
    ('重点关注三：6G标准战', '6G不仅是通信升级，是国家竞争力基础设施，2027-2028年6G标准初步定型'),
    ('重点关注四：AI数据版权战', '将倒逼行业建立数据授权机制，未来1-2年内将出现判例法或立法，明确AI训练数据边界'),
]

for title_text, content in analysis_items:
    p = doc.add_paragraph()
    run = p.add_run(f'■ {title_text}\n')
    run.bold = True
    p.add_run(f'{content}')

doc.add_paragraph()

# Section 5: Conclusion
h5 = doc.add_heading('五、本周核心结论', level=1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '领域'
hdr_cells[1].text = '评级'
hdr_cells[2].text = '核心结论'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

# Data rows
data = [
    ('国际政治', '高风险', '伊朗问题最需关注，中东可能生变'),
    ('AI', '加速', '军事化、版权战、多模态三线并进'),
    ('通信', '稳健', '5G深化落地，6G备战启动'),
]
for i, (area, rating, conclusion) in enumerate(data):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = area
    row_cells[1].text = rating
    row_cells[2].text = conclusion

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('本报告由诸葛亮（首席战略顾问）编制，仅供参考，不构成投资建议\n生成时间：2026年5月11日 北京时间上午8:00')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save('/root/.openclaw/workspace/zhugeliang/每周简报/2026-05-11.docx')
print("Word document created successfully!")
